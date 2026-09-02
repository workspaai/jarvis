"""Streamlit UI dla Jarvisa: /ask (RAG, Week 2), /ingest oraz AGENT (Week 3).

UI tylko woła API albo kod agenta — cała logika (RAG, pętla Think/Act/Observe)
mieszka w FastAPI i modułach `agent_raw` / `agent_graph`. Zero sekretów po
stronie strony: adres API z paska bocznego.

Upload plików (txt/md/docx/xlsx/csv/pdf) parsuje się TUTAJ, w Streamlicie —
do `/ingest` idzie już czysty tekst, więc backend pozostaje bez zmian.

Uwaga: zakładka Agent uruchamia agenta LOKALNIE (import modułu), bo agent nie
jest wystawiony na Renderze — zadanie Week 3 wymaga zrzutu/klipu UI, nie URL-a.

Run:
  streamlit run demo_page.py
"""

import csv
import io
import json
import re
from pathlib import Path

import httpx
import streamlit as st

WORKDIR_CMD = "projekty/jarvis"
MODELS = ["gpt-4o-mini", "gpt-4o", "o3-mini"]
UPLOAD_TYPES = ["txt", "md", "docx", "xlsx", "csv", "pdf"]
PREVIEW_CHARS = 500

# Przełącznik językowy (demo 1:1): etykiety UI + pole `language` w payload /ask —
# wybór promptu PL/EN robi BACKEND. Tłumaczymy rdzeń demo (zakładki /ask i
# /ingest); Agent/Evale/Pamięć zostają po polsku, bo ich treść (ślady, raporty)
# i tak jest polska. Komunikaty błędów parsera plików również zostają po polsku.
UI_TEXTS = {
    "pl": {
        "title": "Jarvis: pytania do dokumentów (RAG)",
        "caption": (
            "Streamlit tylko woła API — cała logika RAG (chunking, retrieval, grounding) "
            "mieszka w FastAPI. Pliki `stages/` pokazują, jak endpoint rósł krok po kroku."
        ),
        "health_button": "Sprawdź /health",
        "tabs": [
            "Zadaj pytanie (/ask)",
            "Dodaj dokument (/ingest)",
            "Agent (Week 3)",
            "Evale (Week 4)",
            "Pamięć (Week 5)",
        ],
        "question_label": "Pytanie",
        "question_placeholder": "Zadaj pytanie o zaindeksowane dokumenty…",
        "model_label": "Model",
        "force_bad_label": "Wymuś zepsutą pierwszą odpowiedź (demo walidacji + retry z Week 1)",
        "reasoned_label": "Rachunek przed werdyktem (naprawa błędnych porównań z liczbami)",
        "ask_button": "Zapytaj",
        "ask_spinner": "Wołam /ask… (uśpiony Render może się budzić do ~1 min)",
        "full_json": "Pełny JSON odpowiedzi",
        "answer_header": "### Odpowiedź",
        "not_in_docs": "Brak w dokumentach",
        "sources_prefix": "**Źródła (cytowane):** ",
        "sources_none": "**Źródła (cytowane):** _brak — odpowiedź nie cytuje żadnego dokumentu_",
        "chunks_prefix": "**Pobrane chunki (retrieval):** ",
        "ingest_caption": "Demo publiczne — guardrail przy ingeście jest w backlogu (Week 3).",
        "docs_header": "📚 Dokumenty w bazie (inwentarz z metadanych — bez LLM i bez retrievalu)",
        "docs_refresh": "Odśwież listę dokumentów",
        "docs_empty": "Baza jest pusta — wgraj pierwszy dokument poniżej.",
        "docs_total": "Razem: **{docs}** dokument(ów), **{chunks}** chunk(ów).",
        "upload_header": "### Wgraj plik",
        "upload_caption": (
            "Parsowanie dzieje się tutaj, w Streamlicie — do API leci już czysty "
            "tekst, więc endpoint `/ingest` pozostaje bez zmian. Obsługiwane: {types}."
        ),
        "upload_label": "Pliki do zaindeksowania",
        "extracted_line": "**{name}** — wyciągnięto **{count}** znaków; podgląd pierwszych {preview}:",
        "docid_label": "document_id (zaproponowany z nazwy pliku — możesz poprawić)",
        "index_file_button": "Zaindeksuj ten plik",
        "docid_required": "Podaj document_id — pole nie może być puste.",
        "manual_header": "### …albo wklej tekst ręcznie",
        "doc_text_label": "Treść dokumentu",
        "doc_text_placeholder": "Wklej pełny tekst dokumentu do zaindeksowania…",
        "docid_manual_placeholder": "np. POL-101",
        "source_label": "source (opcjonalne)",
        "source_placeholder": "np. doc1_handbook.txt",
        "index_button": "Zaindeksuj",
        "ingest_spinner": "Wołam /ingest…",
        "ingest_success": "Zaindeksowano **{doc}** → {chunks} chunk(ów), status: {status}",
    },
    "en": {
        "title": "Jarvis: ask your documents (RAG)",
        "caption": (
            "Streamlit only calls the API — all RAG logic (chunking, retrieval, grounding) "
            "lives in FastAPI. The `stages/` files show how the endpoint grew step by step."
        ),
        "health_button": "Check /health",
        "tabs": [
            "Ask a question (/ask)",
            "Add a document (/ingest)",
            "Agent (Week 3)",
            "Evals (Week 4)",
            "Memory (Week 5)",
        ],
        "question_label": "Question",
        "question_placeholder": "Ask about the indexed documents…",
        "model_label": "Model",
        "force_bad_label": "Force a broken first answer (validation + retry demo from Week 1)",
        "reasoned_label": "Calculation before verdict (numeric-comparison bias fix)",
        "ask_button": "Ask",
        "ask_spinner": "Calling /ask… (a sleeping Render instance may take up to ~1 min)",
        "full_json": "Full JSON response",
        "answer_header": "### Answer",
        "not_in_docs": "Not in the documents",
        "sources_prefix": "**Sources (cited):** ",
        "sources_none": "**Sources (cited):** _none — the answer cites no document_",
        "chunks_prefix": "**Retrieved chunks:** ",
        "ingest_caption": "Public demo — an ingest guardrail is in the backlog (Week 3).",
        "docs_header": "📚 Documents in the index (metadata inventory — no LLM, no retrieval)",
        "docs_refresh": "Refresh document list",
        "docs_empty": "The index is empty — upload your first document below.",
        "docs_total": "Total: **{docs}** document(s), **{chunks}** chunk(s).",
        "upload_header": "### Upload a file",
        "upload_caption": (
            "Parsing happens right here, in Streamlit — the API receives plain "
            "text, so the `/ingest` endpoint stays unchanged. Supported: {types}."
        ),
        "upload_label": "Files to index",
        "extracted_line": "**{name}** — extracted **{count}** characters; preview of the first {preview}:",
        "docid_label": "document_id (proposed from the file name — feel free to edit)",
        "index_file_button": "Index this file",
        "docid_required": "document_id cannot be empty.",
        "manual_header": "### …or paste text manually",
        "doc_text_label": "Document text",
        "doc_text_placeholder": "Paste the full text of the document to index…",
        "docid_manual_placeholder": "e.g. POL-101",
        "source_label": "source (optional)",
        "source_placeholder": "e.g. doc1_handbook.txt",
        "index_button": "Index",
        "ingest_spinner": "Calling /ingest…",
        "ingest_success": "Indexed **{doc}** → {chunks} chunk(s), status: {status}",
    },
}


def build_payload(
    question: str, model: str, force_bad: bool, language: str, reasoned: bool
) -> dict:
    return {
        "question": question,
        "model": model,
        "force_bad": force_bad,
        "language": language,
        "reasoned": reasoned,
    }


def render_curl(base_url: str, payload: dict) -> str:
    body = json.dumps(payload)
    return (
        f'curl -s -X POST {base_url.rstrip("/")}/ask '
        f'-H "Content-Type: application/json" '
        f"-d '{body}'"
    )


def call_json(method: str, url: str, payload: dict | None = None) -> tuple[int, dict | str]:
    try:
        if method == "POST":
            response = httpx.post(url, json=payload, timeout=120.0)
        else:
            response = httpx.get(url, timeout=120.0)

        try:
            return response.status_code, response.json()
        except json.JSONDecodeError:
            return response.status_code, response.text
    except httpx.ConnectError:
        return 0, {"error": f"Cannot reach {url}. Start the API server first."}
    except httpx.HTTPError as exc:
        return 0, {"error": str(exc)}


# --- Parsowanie wgranych plików: dzieje się w CAŁOŚCI w Streamlicie ---------
# Do /ingest leci już czysty tekst — endpoint i reszta backendu bez zmian.


def propose_document_id(filename: str) -> str:
    """„umowa-najmu.docx" → „UMOWA-NAJMU" — propozycja do ręcznej poprawki."""
    stem = Path(filename).stem
    cleaned = re.sub(r"[\s_]+", "-", stem.strip())
    cleaned = re.sub(r"[^\w-]", "", cleaned)
    cleaned = re.sub(r"-{2,}", "-", cleaned).strip("-")
    return cleaned.upper() or "DOKUMENT"


def _decode_text(raw: bytes) -> str:
    # UTF-8 (także z BOM) pierwszy; CP1250 łapie pliki z polskiego Windowsa.
    for encoding in ("utf-8-sig", "cp1250"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError(
        "nie udało się odczytać pliku jako tekstu (próbowałem UTF-8 i CP1250) — "
        "to raczej plik binarny albo w nietypowym kodowaniu."
    )


def _extract_csv(raw: bytes) -> str:
    text = _decode_text(raw)
    try:
        dialect = csv.Sniffer().sniff(text[:2048], delimiters=",;\t")
    except csv.Error:
        dialect = csv.excel  # nie wykrył separatora → przecinek jak w Excelu
    rows = csv.reader(io.StringIO(text), dialect)
    # Komórki sklejamy przez „ | ": chunk czyta się jak tabela, nie jak zlepek.
    lines = [
        " | ".join(cell.strip() for cell in row)
        for row in rows
        if any(cell.strip() for cell in row)
    ]
    return "\n".join(lines)


def _extract_docx(raw: bytes) -> str:
    # Import w środku — strona ma wstać także bez doinstalowanych parserów.
    from docx import Document

    document = Document(io.BytesIO(raw))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _extract_xlsx(raw: bytes) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    parts = []
    for sheet in workbook.worksheets:
        parts.append(f"[Arkusz: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(value).strip() for value in row if value is not None]
            if cells:
                parts.append(" | ".join(cells))
    workbook.close()
    return "\n".join(parts)


def _extract_pdf(raw: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    if reader.is_encrypted:
        try:
            reader.decrypt("")  # puste hasło otwiera PDF-y „zamknięte" tylko technicznie
        except Exception:
            raise ValueError("PDF jest zabezpieczony hasłem — wgraj wersję bez hasła.")
    text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    if not text:
        raise ValueError(
            f"PDF nie ma warstwy tekstowej ({len(reader.pages)} str.) — to "
            "prawdopodobnie skan z obrazów. Potrzebny OCR, albo wklej tekst "
            "ręcznie w polu poniżej."
        )
    return text


@st.cache_data(show_spinner=False)
def extract_text_from_upload(filename: str, raw: bytes) -> str:
    """Czysty tekst z pliku; ValueError z polskim komunikatem, gdy się nie da."""
    suffix = Path(filename).suffix.lower()
    extractors = {
        ".txt": _decode_text,
        ".md": _decode_text,
        ".csv": _extract_csv,
        ".docx": _extract_docx,
        ".xlsx": _extract_xlsx,
        ".pdf": _extract_pdf,
    }
    if suffix not in extractors:
        raise ValueError(f"nieobsługiwane rozszerzenie: {suffix or 'brak'}.")

    try:
        text = extractors[suffix](raw)
    except ValueError:
        raise  # nasze komunikaty przechodzą bez opakowania
    except ImportError as exc:
        raise ValueError(
            f"brak biblioteki do tego formatu ({exc.name}) — uruchom: "
            "pip install -r requirements.txt."
        ) from exc
    except Exception as exc:
        raise ValueError(
            f"nie udało się odczytać pliku ({type(exc).__name__}) — plik może "
            "być uszkodzony albo w innym formacie, niż wskazuje rozszerzenie."
        ) from exc

    text = text.strip()
    if not text:
        raise ValueError("plik nie zawiera żadnego tekstu — nie ma czego indeksować.")
    return text


def ingest_and_render(
    base_url: str, text: str, document_id: str, source: str, texts: dict
) -> None:
    """Wysyła tekst do /ingest i rysuje wynik — wspólne dla uploadu i wklejki."""
    payload = {"text": text, "document_id": document_id, "source": source}
    with st.spinner(texts["ingest_spinner"]):
        status, data = call_json("POST", f"{base_url.rstrip('/')}/ingest", payload)
    st.markdown(f"**HTTP {status}**" if status else "**Request failed**")
    if status == 200 and isinstance(data, dict):
        st.success(
            texts["ingest_success"].format(
                doc=data.get("document_id"),
                chunks=data.get("chunks_indexed"),
                status=data.get("status"),
            )
        )
    st.json(data)


def render_attempts(data: dict | str) -> None:
    if not isinstance(data, dict):
        return

    attempts = data.get("attempts", [])
    if not attempts:
        return

    st.markdown("### Attempts")
    for attempt in attempts:
        status = "passed" if attempt.get("ok") else "failed"
        title = f"Attempt {attempt.get('attempt')}: {attempt.get('step')} ({status})"
        with st.expander(title, expanded=True):
            st.write(attempt.get("message"))
            if attempt.get("raw_output"):
                st.markdown("**Raw model output**")
                st.code(attempt["raw_output"], language="json")
            if attempt.get("validation_error"):
                st.markdown("**Validation error**")
                st.code(attempt["validation_error"], language="text")


def render_response_summary(data: dict | str, texts: dict) -> None:
    if not isinstance(data, dict) or "error" in data:
        return

    answer = data.get("answer")
    if isinstance(answer, dict):
        st.markdown(texts["answer_header"])
        if answer.get("i_dont_know"):
            # Odmowa ma być czytelna od razu, nie ukryta w polu boolean.
            st.warning(f"**{texts['not_in_docs']}** — {answer.get('answer', '')}")
        else:
            st.success(answer.get("answer", ""))

        source = answer.get("source") or []
        if source:
            st.markdown(texts["sources_prefix"] + " ".join(f"`{s}`" for s in source))
        else:
            st.markdown(texts["sources_none"])
        st.caption(
            f"confidence: {answer.get('confidence')} | "
            f"sources_needed: {answer.get('sources_needed')} | "
            f"i_dont_know: {answer.get('i_dont_know')}"
        )

    retrieved = data.get("retrieved_chunks") or []
    if retrieved:
        st.markdown(texts["chunks_prefix"] + " ".join(f"`{c}`" for c in retrieved))

    metric_cols = st.columns(4)
    metric_cols[0].metric("Model", str(data.get("model", "-")))
    metric_cols[1].metric("Tokens", str(data.get("tokens_used", "-")))
    metric_cols[2].metric("Latency", f"{data.get('latency_ms', '-')} ms")
    metric_cols[3].metric("Cost", f"${data.get('cost_usd', '-')}")


def render_eval_report(eval_report: dict, saved: bool) -> None:
    """Rysuje raport evali; wołający musi łapać wyjątki (patrz zakładka Evale)."""
    summary = eval_report["summary"]
    baseline = None
    baseline_path = Path(__file__).resolve().parent / "evals" / "report_baseline.json"
    if baseline_path.exists():
        baseline = json.loads(baseline_path.read_text(encoding="utf-8"))

    st.markdown("### Wynik zestawu")
    if saved:
        st.info(
            "Na tej instancji nie ma kompletu trace'ów (katalog `traces/` nie "
            "wchodzi do repo) — to ZAPISANY wynik ostatniego pełnego przebiegu "
            "lokalnego (`evals/report_after_fix3.json`), nie świeże liczenie."
        )
    delta = None
    if baseline:
        delta_pp = (summary["pass_rate"] - baseline["summary"]["pass_rate"]) * 100
        delta = f"{delta_pp:+.1f} pp vs baseline"
    eval_cols = st.columns(3)
    eval_cols[0].metric("PASS RATE", f"{summary['pass_rate']:.1%}", delta)
    eval_cols[1].metric(
        "Asercje zaliczone", f"{summary['checks_passed']}/{summary['checks_total']}"
    )
    eval_cols[2].metric("Pytań ocenionych", summary["questions"])

    st.markdown("### Per asercja — baseline (PRZED) vs teraz (PO)")
    table_rows = []
    for name in sorted(eval_report["per_assertion"]):
        now = eval_report["per_assertion"][name]
        row = {"asercja": name}
        if baseline and name in baseline["per_assertion"]:
            b = baseline["per_assertion"][name]
            row["baseline"] = f"{b['pass']}/{b['total']}"
        row["teraz"] = f"{now['pass']}/{now['total']}"
        if baseline and name in baseline["per_assertion"]:
            row["zmiana"] = now["pass"] - baseline["per_assertion"][name]["pass"]
        table_rows.append(row)
    st.dataframe(table_rows, width="stretch", hide_index=True)

    if eval_report["failures"]:
        st.markdown("### Porażki (werdykt binarny + powód jednym zdaniem)")
        for failure in eval_report["failures"]:
            st.warning(
                f"**{failure['qid']}** · {failure['engine']} · "
                f"`{failure['assertion']}` — {failure['reason']}"
            )
    else:
        st.success("Zero porażek.")

    st.caption(
        "Baseline = stan sprzed poprawek promptu (evals/report_baseline.json). "
        "Known-fail: q12 — fałszywa odmowa przy pytaniu wymagającym wnioskowania "
        "(3 nieudane iteracje promptowe; kandydat na mocniejszy model w tej roli)."
    )


st.set_page_config(page_title="Jarvis — demo RAG (Week 2)", layout="centered")

# Przełącznik języka PRZED tytułem — etykiety całej strony zależą od wyboru.
# Uwaga demo: zmiana języka przebudowuje zakładki (wracają do pierwszej),
# więc język wybieramy RAZ, na początku pokazu.
language_choice = st.sidebar.radio(
    "Język / Language", ["Polski", "English"], horizontal=True, key="ui_language"
)
LANG = "pl" if language_choice == "Polski" else "en"
T = UI_TEXTS[LANG]

st.title(T["title"])
st.caption(T["caption"])

base_url = st.sidebar.text_input("API base URL", "https://jarvis-8lpg.onrender.com")
st.sidebar.markdown("### Start the API")
st.sidebar.code(
    f"cd {WORKDIR_CMD}\n"
    "source .venv/bin/activate\n"
    "uvicorn main:app --host 127.0.0.1 --port 8000 --reload",
    language="bash",
)
st.sidebar.markdown("### Start this page")
st.sidebar.code(
    f"cd {WORKDIR_CMD}\nsource .venv/bin/activate\nstreamlit run demo_page.py",
    language="bash",
)

if st.sidebar.button(T["health_button"]):
    status, data = call_json("GET", f"{base_url.rstrip('/')}/health")
    st.sidebar.markdown(f"**HTTP {status}**" if status else "**Brak połączenia**")
    st.sidebar.json(data)

tab_ask, tab_ingest, tab_agent, tab_evals, tab_memory = st.tabs(T["tabs"])

with tab_ask:
    with st.form("ask_form"):
        question = st.text_area(
            T["question_label"],
            height=100,
            placeholder=T["question_placeholder"],
        )
        model = st.selectbox(T["model_label"], MODELS, index=0)
        force_bad = st.checkbox(T["force_bad_label"], value=False)
        # Domyślnie WŁĄCZONE w UI (naprawa Q3); API bez flagi zostaje przy starej
        # ścieżce, więc evale i dotychczasowe klienty są nietknięte.
        reasoned = st.checkbox(T["reasoned_label"], value=True)
        submitted = st.form_submit_button(T["ask_button"], type="primary")

    payload = build_payload(question, model, force_bad, LANG, reasoned)
    st.markdown("#### Request")
    st.code(render_curl(base_url, payload), language="bash")

    if submitted:
        with st.spinner(T["ask_spinner"]):
            status, data = call_json("POST", f"{base_url.rstrip('/')}/ask", payload)
        st.markdown(f"**HTTP {status}**" if status else "**Request failed**")
        render_response_summary(data, T)
        render_attempts(data)
        with st.expander(T["full_json"]):
            st.json(data)

with tab_ingest:
    st.caption(T["ingest_caption"])

    # Inwentarz z metadanych (GET /documents) — odpowiedź na klasę „wymień
    # wszystkie", której retrieval semantyczny nie pokrywa (odkrycie W5/W6).
    with st.expander(T["docs_header"]):
        if st.button(T["docs_refresh"], key="docs_refresh_btn"):
            status, data = call_json("GET", f"{base_url.rstrip('/')}/documents")
            if status == 200 and isinstance(data, dict):
                if data.get("documents"):
                    for doc in data["documents"]:
                        st.markdown(f"- **{doc['document_id']}** — {doc['chunks']} chunk(ów)")
                    st.caption(
                        T["docs_total"].format(
                            docs=len(data["documents"]), chunks=data.get("total_chunks", 0)
                        )
                    )
                else:
                    st.info(T["docs_empty"])
            else:
                st.markdown(f"**HTTP {status}**" if status else "**Request failed**")
                st.json(data)

    st.markdown(T["upload_header"])
    st.caption(
        T["upload_caption"].format(
            types=", ".join(f".{ext}" for ext in UPLOAD_TYPES)
        )
    )
    uploads = st.file_uploader(
        T["upload_label"],
        type=UPLOAD_TYPES,
        accept_multiple_files=True,
        key="ingest_uploads",
    )
    for position, upload in enumerate(uploads or []):
        if position:
            st.divider()
        try:
            extracted = extract_text_from_upload(upload.name, upload.getvalue())
        except ValueError as exc:
            # Zły plik ≠ wyjątek na ekranie: pokazujemy powód i idziemy dalej.
            st.error(f"**{upload.name}** — {exc}")
            continue

        st.markdown(
            T["extracted_line"].format(
                name=upload.name,
                count=len(extracted),
                preview=min(len(extracted), PREVIEW_CHARS),
            )
        )
        preview_suffix = "…" if len(extracted) > PREVIEW_CHARS else ""
        st.code(
            extracted[:PREVIEW_CHARS] + preview_suffix,
            language=None,
            wrap_lines=True,
            height=200,
        )
        upload_document_id = st.text_input(
            T["docid_label"],
            value=propose_document_id(upload.name),
            key=f"upload_docid_{position}_{upload.name}",
        )
        if st.button(
            T["index_file_button"],
            key=f"upload_send_{position}_{upload.name}",
            type="primary",
        ):
            if not upload_document_id.strip():
                st.warning(T["docid_required"])
            else:
                ingest_and_render(
                    base_url, extracted, upload_document_id.strip(), upload.name, T
                )

    st.markdown(T["manual_header"])
    with st.form("ingest_form"):
        doc_text = st.text_area(
            T["doc_text_label"],
            height=220,
            placeholder=T["doc_text_placeholder"],
        )
        document_id = st.text_input(
            "document_id", placeholder=T["docid_manual_placeholder"]
        )
        doc_source = st.text_input(
            T["source_label"], placeholder=T["source_placeholder"]
        )
        ingest_submitted = st.form_submit_button(T["index_button"], type="primary")

    if ingest_submitted:
        ingest_and_render(base_url, doc_text, document_id, doc_source, T)

with tab_agent:
    st.markdown(
        "Agent **planuje → woła narzędzie → ocenia wynik → decyduje ponownie**. "
        "Drugie wyszukiwanie następuje TYLKO wtedy, gdy pierwsze nie odpowiedziało "
        "na pytanie — dlatego to agent, a nie sztywny workflow."
    )
    st.caption(
        "Agent działa lokalnie (import modułu `agent_raw` / `agent_graph`), "
        "korzysta z tej samej bazy wektorowej co /ask. UI nie zawiera logiki agenta."
    )
    st.info(
        "Agent działa na lokalnej bazie wektorowej — na wdrożonej instancji "
        "korpus jest pusty. Dowód działania: zrzuty w zgłoszeniu Week 3."
    )

    # Formularz: przełącznik silnika i pytanie NIE wywołują przeładowania strony
    # (Streamlit przy każdym rerunie wracał do pierwszej zakładki — psuło demo).
    with st.form("agent_form"):
        agent_question = st.text_area(
            "Pytanie do agenta",
            key="agent_question",
            height=80,
            placeholder="Zapytaj o coś z zaindeksowanych dokumentów…",
        )
        engine_label = st.radio(
            "Silnik pętli",
            ["Surowa pętla (agent_raw)", "LangGraph StateGraph (agent_graph)"],
            horizontal=True,
            help="Ta sama logika, dwie maszynerie — dowód z Kroku 3 dziennika budowy.",
        )
        run_clicked = st.form_submit_button("Uruchom agenta", type="primary")

    if run_clicked:
        if not (agent_question or "").strip():
            st.warning("Wpisz pytanie do agenta.")
        else:
            # Import w środku: Streamlit ma wystartować także wtedy, gdy ktoś
            # używa tylko zakładek /ask i /ingest (np. bez klucza OpenAI).
            if engine_label.startswith("Surowa"):
                from agent_raw import run_agent as run_agent_fn

                engine_name = "surowa pętla (agent_raw.py)"
                kwargs = {}
            else:
                from agent_graph import run_agent as run_agent_fn

                engine_name = "LangGraph StateGraph (agent_graph.py)"
                # Bez stałego thread_id: każde pytanie ma dostać świeży stan
                # (checkpointer + reducery dokleiłyby poprzedni przebieg).
                kwargs = {}

            with st.spinner(f"Agent pracuje — {engine_name}…"):
                try:
                    result = run_agent_fn(agent_question, **kwargs)
                except Exception as exc:  # błąd pokazujemy, nie chowamy
                    st.error(f"Agent nie dokończył pracy: {exc}")
                    result = None

            # Wynik do session_state: Streamlit przeładowuje skrypt przy KAŻDEJ
            # zmianie rozmiaru okna albo powiększenia, a sekcja renderowana tylko
            # pod `run_clicked` znikała wtedy z ekranu (psuło robienie zrzutów).
            if result:
                st.session_state["agent_result"] = result
                st.session_state["agent_engine"] = engine_name

    result = st.session_state.get("agent_result")
    engine_name = st.session_state.get("agent_engine", "")
    if result:
        st.markdown("### Odpowiedź agenta")
        if result["refused"]:
            st.warning(f"**Uczciwa odmowa** — {result['answer']}")
        else:
            st.success(result["answer"])

        if result["sources"]:
            st.markdown(
                "**Źródła (cytowane):** "
                + " ".join(f"`{s}`" for s in result["sources"])
            )
        else:
            st.markdown(
                "**Źródła (cytowane):** _brak — agent nie znalazł podstaw_"
            )

        st.markdown("### Ślad pętli agenta (Think → Act → Observe)")
        st.caption(
            "ACT pokazuje, JAKIEJ FRAZY agent szukał w danym kroku — "
            "przy nieudanym wyszukiwaniu sam ją przeformułowuje."
        )
        current_iteration = None
        for step in result["trace"]:
            if step["iteration"] != current_iteration:
                current_iteration = step["iteration"]
                st.markdown(f"**— iteracja {current_iteration} —**")
            icon = {
                "THINK": "🧠 THINK",
                "ACT": "🔧 ACT",
                "OBSERVE": "👁 OBSERVE",
                "STOP": "🛑 STOP",
            }.get(step["kind"], step["kind"])
            st.markdown(f"- {icon} — `{step['text']}`")

        st.markdown("### Metryki przebiegu")
        # Cztery kolumny, a koszt w podpisie: przy pięciu kolumnach
        # Streamlit ucinał wartość do "$0.00…".
        agent_cols = st.columns(4)
        agent_cols[0].metric("Iteracje", result["iterations"])
        agent_cols[1].metric("Wywołania narzędzia", result["tool_calls"])
        agent_cols[2].metric("Tokeny", result["tokens"])
        agent_cols[3].metric("Czas", f"{result['latency_ms']} ms")
        st.caption(
            f"Koszt przebiegu: **${result['cost_usd']:.6f}** · "
            f"silnik: {engine_name} · limit iteracji: 6 (fail closed)"
        )

with tab_evals:
    st.markdown(
        "Zestaw **asercji kodowych** (Codify z pętli TRACE): zero wywołań LLM — "
        "czyta zapisane trace'y z `traces/traces.jsonl` i ocenia **najnowszy** "
        "przebieg każdego z 20 pytań. Każda asercja jest powiązana z kategorią "
        "z taksonomii awarii zbudowanej przy open codingu."
    )
    st.caption(
        "TRACE w praktyce: Trace (zapis przebiegów) → Read/Analyze (open coding "
        "i taksonomia) → Codify (te asercje) → Enforce (ten przycisk; raporty "
        "przed/po w evals/)."
    )

    if st.button("Uruchom zestaw evali", type="primary", key="run_evals_btn"):
        # Import w środku — zakładki /ask i /ingest mają działać także bez
        # wygenerowanych trace'ów (np. świeży klon repo).
        from evals.run_evals import KIND, build_report

        with st.spinner("Liczę asercje na najnowszych trace'ach…"):
            try:
                try:
                    report = build_report()
                except FileNotFoundError:
                    report = None  # instancja bez traces/ (katalog nie wchodzi do repo)
                # Świeży wynik tylko z KOMPLETU pytań: na wdrożonej instancji
                # traces.jsonl ma co najwyżej pojedyncze przebiegi z zakładki
                # Agent — liczenie z nich udawałoby pełny eval.
                if report and report["summary"]["questions"] == len(KIND):
                    st.session_state["eval_report"] = report
                    st.session_state["eval_report_saved"] = False
                else:
                    saved_path = (
                        Path(__file__).resolve().parent / "evals" / "report_after_fix3.json"
                    )
                    st.session_state["eval_report"] = json.loads(
                        saved_path.read_text(encoding="utf-8")
                    )
                    st.session_state["eval_report_saved"] = True
            except Exception as exc:  # pokazujemy wprost, nie chowamy
                st.session_state.pop("eval_report", None)
                st.error(f"Nie udało się policzyć evali: {exc}")

    eval_report = st.session_state.get("eval_report")
    if eval_report:
        try:
            render_eval_report(
                eval_report, st.session_state.get("eval_report_saved", False)
            )
        except Exception as exc:
            # Przełączenie zakładki = rerun CAŁEGO skryptu, więc raport, którego
            # nie da się narysować, ubijałby wszystkie zakładki — leci ze stanu.
            st.session_state.pop("eval_report", None)
            st.error(f"Nie udało się wyświetlić raportu evali: {exc}")

with tab_memory:
    st.markdown(
        "**Pamięć długoterminowa Jarvisa (Week 5).** Fakty żyją w ZEWNĘTRZNEJ bazie "
        "(Neon Postgres + pgvector), więc **przeżywają restart procesu i nową sesję** — "
        "to NIE historia czatu, tylko trwały store. Zapis przechodzi przez BRAMKĘ "
        "(reguły + mały model decydują, czy to trwały fakt); odczyt jest hybrydowy "
        "(podobieństwo znaczeniowe + świeżość)."
    )

    st.markdown("### 1) Zapisz preferencję / fakt")
    with st.form("mem_save_form"):
        save_text = st.text_input(
            "Powiedz Jarvisowi coś do zapamiętania",
            placeholder="np. Zapisz, że rozliczam projekty w EUR",
        )
        save_clicked = st.form_submit_button("Zapisz do pamięci", type="primary")

    if save_clicked:
        if not (save_text or "").strip():
            st.warning("Wpisz coś do zapamiętania.")
        else:
            with st.spinner("Bramka zapisu ocenia turę…"):
                try:
                    # Import w środku — zakładki /ask i /ingest mają działać także bez
                    # skonfigurowanej bazy (DATABASE_URL) czy klucza OpenAI.
                    from write_gate import maybe_write_memory

                    saved = maybe_write_memory(save_text, source_event_id="streamlit-demo")
                    if saved:
                        st.success(
                            f"✅ Zapisano trwale: **{saved['fact']}**  ·  "
                            f"_(temat: {saved['subject']})_"
                        )
                    else:
                        st.info(
                            "Bramka NIE zapisała — to nie był trwały fakt/preferencja "
                            "(np. pytanie albo treść przejściowa). Tak ma działać: "
                            "store zostaje mały i trafny."
                        )
                except Exception as exc:  # błąd pokazujemy wprost, nie chowamy
                    st.error(f"Błąd zapisu: {exc}")

    st.divider()

    st.markdown("### 2) Odczyt w NOWEJ sesji")
    st.caption(
        "Odczyt idzie PROSTO do bazy Neon — nie z pamięci tej sesji. Dowód recall między "
        "sesjami: otwórz ten URL w nowej karcie / oknie incognito i kliknij tylko odczyt "
        "— fakt nadal wróci, bez powtarzania go."
    )
    with st.form("mem_recall_form"):
        recall_text = st.text_input(
            "Zapytaj o coś, co Jarvis mógł zapamiętać",
            placeholder="np. w czym rozliczam projekty?",
        )
        recall_clicked = st.form_submit_button("Odczytaj z pamięci", type="primary")

    if recall_clicked:
        if not (recall_text or "").strip():
            st.warning("Wpisz pytanie.")
        else:
            with st.spinner("Szukam w pamięci (podobieństwo + świeżość)…"):
                try:
                    from memory_store import recall_facts

                    hits = recall_facts(recall_text, limit=3)
                    if hits:
                        st.markdown("**Jarvis pamięta:**")
                        for h in hits:
                            st.success(
                                f"🧠 {h[2]}  ·  _(temat: {h[1]}, trafność {h[4]:.2f})_"
                            )
                    else:
                        st.info("Nic nie znalazłem w pamięci dla tego pytania.")
                except Exception as exc:
                    st.error(f"Błąd odczytu: {exc}")

    with st.expander("Podgląd całego trwałego store'u (co przetrwało między sesjami)"):
        try:
            from memory_store import get_facts

            all_facts = get_facts()
            if all_facts:
                st.dataframe(
                    [
                        {
                            "id": f[0],
                            "temat": f[1],
                            "fakt": f[2],
                            "zapisano": str(f[4])[:19],
                        }
                        for f in all_facts
                    ],
                    hide_index=True,
                    width="stretch",
                )
            else:
                st.caption("Store jest pusty — zapisz pierwszy fakt powyżej.")
        except Exception as exc:
            st.error(f"Nie udało się odczytać store'u: {exc}")
